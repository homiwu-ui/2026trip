from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

ITALIAN_GREEN = RGBColor(0x00, 0x92, 0x46)
ITALIAN_RED = RGBColor(0xCF, 0x2B, 0x2B)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ITALIAN_GREEN
    return heading

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_schedule_table(doc, schedule, day_title):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['時間', '活動', '備註']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "009246")

    for idx, item in enumerate(schedule):
        row = table.add_row()
        row.cells[0].text = item['time']
        row.cells[1].text = item['activity']
        row.cells[2].text = item['note']
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F0F7F0")

    for row in table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(7.5)
        row.cells[2].width = Cm(5.5)

def add_attractions(doc, attractions):
    add_heading_styled(doc, "景點介紹", level=3)
    for attr in attractions:
        p = doc.add_paragraph()
        run_name = p.add_run(f"📍 {attr['name']}")
        run_name.bold = True
        run_name.font.size = Pt(10)
        run_name.font.color.rgb = ITALIAN_RED

        p2 = doc.add_paragraph(attr['description'])
        for run in p2.runs:
            run.font.size = Pt(9.5)

        p3 = doc.add_paragraph()
        run_tip_label = p3.add_run("💡 Tips：")
        run_tip_label.bold = True
        run_tip_label.font.size = Pt(9)
        run_tip_label.font.color.rgb = ITALIAN_GREEN
        run_tip = p3.add_run(attr['tip'])
        run_tip.font.size = Pt(9)

def add_food_list(doc, food_list):
    add_heading_styled(doc, "餐飲推薦", level=3)
    for f in food_list:
        p = doc.add_paragraph()
        run_name = p.add_run(f"{f['name']}  ")
        run_name.bold = True
        run_name.font.size = Pt(10)
        run_info = p.add_run(f"({f['type']})　推薦 {f['recommend']}　{f['price']}")
        run_info.font.size = Pt(9)
        run_info.font.color.rgb = DARK_GRAY

def add_transport(doc, transport):
    add_heading_styled(doc, "交通方式", level=3)
    for t in transport:
        p = doc.add_paragraph()
        run = p.add_run(f"{t['mode']}  {t['detail']}")
        run.font.size = Pt(9.5)

def add_tips(doc, tips):
    add_heading_styled(doc, "小提醒", level=3)
    for tip in tips:
        p = doc.add_paragraph()
        run = p.add_run(f"• {tip}")
        run.font.size = Pt(9)

# ===== Day 1 =====
add_heading_styled(doc, "Day 1 — 2026年7月22日（週三）", level=1)
p = doc.add_paragraph()
run = p.add_run("羅馬 — 競技場初體驗")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = ITALIAN_RED

add_heading_styled(doc, "行程時間表", level=2)
day1_schedule = [
    {"time": "06:25 - 07:45", "activity": "抵達 Fiumicino 機場 ➔ 入境領行李", "note": "預留 1.5 小時通關"},
    {"time": "08:00 - 09:00", "activity": "Leonardo Express ➔ Roma Termini", "note": "€14/人，32 分鐘直達"},
    {"time": "09:00 - 10:00", "activity": "🧳 寄放行李 @ Termini Deposito Bagagli", "note": "月台 24 旁，€6/件"},
    {"time": "10:00 - 11:30", "activity": "🥐 早午餐 @ Mercato Centrale Roma", "note": "Bonci Pizza al taglio 🦐 Tre Rotelle 96 分！"},
    {"time": "11:30 - 12:30", "activity": "Termini 周邊自由活動／休息", "note": "車站有冷氣"},
    {"time": "12:30 - 13:00", "activity": "🚇 地鐵 B 線 Termini ➔ Colosseo", "note": "2 站"},
    {"time": "13:10 - 15:00", "activity": "🏟️ 羅馬競技場（已預約）", "note": "含地下層參觀"},
    {"time": "15:00 - 16:00", "activity": "古羅馬廣場 & 帕拉提諾山", "note": "聯票包含"},
    {"time": "16:00 - 16:30", "activity": "聖克萊門特教堂（三層時光機）", "note": "€10，競技場步行 3 分"},
    {"time": "16:30 - 17:00", "activity": "🚇 Colosseo ➔ Termini 取行李", "note": "1 站"},
    {"time": "17:00 - 18:00", "activity": "🚇 A 線 Termini ➔ Spagna ➔ Check-in", "note": "2 站至飯店"},
    {"time": "18:00 - 19:30", "activity": "🍝 晚餐 @ Pastificio Guerra 🦐", "note": "站立式 €5 新鮮生麵"},
    {"time": "19:30 - 20:00", "activity": "🍨 Gelato @ La Strega Nocciola", "note": "榛果與開心果必點"},
    {"time": "20:00~", "activity": "超市採買 & 回飯店休息", "note": "第一天別太晚睡"},
]
add_schedule_table(doc, day1_schedule, "Day 1")

add_transport(doc, [
    {"mode": "✈️ 飛機", "detail": "桃園機場 ➔ 羅馬 Fiumicino 機場 (T3)"},
    {"mode": "🚆 機場快線", "detail": "Leonardo Express：Fiumicino ➔ Termini，€14，32 分鐘直達"},
    {"mode": "🚇 地鐵 B", "detail": "Termini ➔ Colosseo，2 站，約 5 分鐘"},
    {"mode": "🚇 地鐵 A", "detail": "Termini ➔ Spagna，2 站（飯店方向）"},
    {"mode": "🚶 步行", "detail": "Colosseo→古羅馬廣場→聖克萊門特，各點步行 3 分鐘"},
])

add_attractions(doc, [
    {
        "name": "羅馬競技場 (Colosseum)",
        "description": "古羅馬帝國最具代表性的建築，曾是角鬥士競技的場所。建議提前在官網預約門票，避免排隊等候。內部可參觀地下層結構與競技場平台。",
        "tip": "建議購買聯票（競技場 + 古羅馬廣場 + 帕拉提諾山）"
    },
    {
        "name": "聖克萊門特教堂 (Basilica di San Clemente)",
        "description": "獨一無二的三層「時光機」教堂——地面層為 12 世紀中世紀教堂 → 地下層為 4 世紀早期基督教會 → 最深層為 1 世紀羅馬密特拉神殿，還能聽到地下水流聲。",
        "tip": "門票 €10，三層全可參觀，小朋友會愛地下探險感"
    }
])

add_food_list(doc, [
    {"name": "Mercato Centrale Roma", "type": "早午餐", "recommend": "Bonci Pizza al taglio（Tre Rotelle 96分）", "price": "€5-10"},
    {"name": "Pastificio Guerra", "type": "晚餐", "recommend": "每日兩種現煮生麵", "price": "約 €5"},
    {"name": "La Strega Nocciola", "type": "Gelato", "recommend": "榛果與開心果口味", "price": "€3-5"},
    {"name": "Pinsitaly Trevi", "type": "備用晚餐", "recommend": "Pinsa 羅馬式橢圓披薩", "price": "€8-12"},
])

add_tips(doc, [
    "穿厚底好走的球鞋，羅馬石板路多",
    "帶空水壺，街頭 Nasone 小噴泉可生飲",
    "後背包請背前面，注意防範扒手",
    "有人靠近要綁幸運繩請堅定說 NO",
    "盡量撐到當地晚上 9 點再睡以調整時差",
    "Termini 寄物處旺季排隊約 20-30 分，預留時間",
])

doc.add_page_break()

# ===== Day 2 =====
add_heading_styled(doc, "Day 2 — 2026年7月23日（週四）", level=1)
p = doc.add_paragraph()
run = p.add_run("羅馬 — 永恆之城深度漫遊")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = ITALIAN_RED

p = doc.add_paragraph()
run = p.add_run("動線：住宿 → 鎖鏈堂 → 萬神殿 → 金杯咖啡 → 耶穌教堂 → Trastevere → 天使教堂 → 西班牙階梯 → 許願池 → 住宿")
run.font.size = Pt(9)
run.font.italic = True

add_heading_styled(doc, "行程時間表", level=2)
day2_schedule = [
    {"time": "08:00 - 08:30", "activity": "早餐 @ 住宿附近", "note": "簡單吃，早點出發"},
    {"time": "08:30 - 09:00", "activity": "步行經鎖鏈堂 (San Pietro in Vincoli)", "note": "朝聖《摩西像》，早上人最少！"},
    {"time": "09:00 - 10:30", "activity": "萬神殿參觀", "note": "上午光線穿過穹頂圓洞最美"},
    {"time": "10:30 - 11:30", "activity": "金杯咖啡 & 納沃納廣場", "note": "Granita al Caffè 必點"},
    {"time": "11:30 - 12:30", "activity": "耶穌教堂 (Church of the Gesù)", "note": "巴洛克幻覺天花板，免費入場"},
    {"time": "12:30 - 14:00", "activity": "午餐 @ Trastevere", "note": "Da Enzo al 29 或 Checco er Carettiere"},
    {"time": "14:00 - 15:00", "activity": "天使與殉教者聖母教堂", "note": "貝尼尼親雕天使像"},
    {"time": "15:00 - 16:30", "activity": "西班牙階梯 & 特雷維許願池", "note": "下午補拍照"},
    {"time": "16:30 - 17:30", "activity": "自由時間 / 購物", "note": "名牌大街或小巷探險"},
    {"time": "18:00 - 19:30", "activity": "晚餐", "note": "品嚐羅馬經典料理"},
    {"time": "20:00 ~", "activity": "夜間散步回住宿", "note": "體驗羅馬夜晚的浪漫"},
]
add_schedule_table(doc, day2_schedule, "Day 2")

add_transport(doc, [
    {"mode": "🚶 步行", "detail": "住宿→鎖鏈堂約 15 分，鎖鏈堂→萬神殿約 20 分"},
    {"mode": "🚶 步行", "detail": "萬神殿→金杯咖啡→耶穌教堂→Trastevere，全程步行 5-15 分"},
    {"mode": "🚶 步行", "detail": "Trastevere→天使教堂→西班牙階梯→許願池，全程步行 10-20 分"},
])

add_attractions(doc, [
    {
        "name": "聖伯多祿鎖鏈堂 (San Pietro in Vincoli)",
        "description": "收藏米開朗基羅曠世巨作《摩西像》——摩西頭上的「角」源自拉丁文聖經翻譯錯誤，成為藝術史上最著名的「美麗錯誤」。清晨人少，可獨享大師雕塑。",
        "tip": "免費入場，早上 8:30 幾乎無人"
    },
    {
        "name": "萬神殿 (Pantheon)",
        "description": "保存最完整的古羅馬建築，其混凝土穹頂至今仍是世界最大的無鋼筋穹頂。內部光線從頂部圓洞灑落，極具神聖感。",
        "tip": "免費入場，上午光線穿過穹頂圓洞最美"
    },
    {
        "name": "耶穌教堂 (Church of the Gesù)",
        "description": "全世界第一座耶穌會教堂，巴洛克風格的鼻祖。天花板壁畫《耶穌之名的勝利》是透視幻覺畫的巔峰——難以分辨何處是畫、何處是雕塑。",
        "tip": "免費入場，抬頭看天花板至少 5 分鐘"
    },
    {
        "name": "天使與殉教者聖母教堂",
        "description": "收藏貝尼尼親手雕刻的天使像，免費入場的藝術寶庫。氣氛神聖寧靜，是避開人潮的好去處。",
        "tip": "這兩座天使像原本要放在聖天使橋上，因太美被收藏在此。"
    },
    {
        "name": "西班牙階梯 (Spanish Steps)",
        "description": "羅馬最經典的電影場景之一。走上階梯中段回頭往下拍，可以拍到名牌大街的繁華視角。注意階梯上禁止飲食與坐下。",
        "tip": "旁邊有羅馬第一家麥當勞，Tiramisu 意外地好吃！"
    },
    {
        "name": "特雷維許願池 (Trevi Fountain)",
        "description": "羅馬最大、最華麗的巴洛克噴泉。建議走到噴泉右側角落，通常人比較少。",
        "tip": "背對噴泉用右手向左肩後方拋硬幣，保證會再回到羅馬！"
    }
])

add_food_list(doc, [
    {"name": "La Casa del Caffè Tazza d'Oro", "type": "咖啡", "recommend": "Granita al Caffè（咖啡冰沙）", "price": "€3-5"},
    {"name": "Da Enrico al 29", "type": "午餐", "recommend": "Carbonara 義大利麵", "price": "€12-18"},
])

add_tips(doc, [
    "鎖鏈堂免費，早點去避開旅行團",
    "萬神殿免費入場，注意穿著（禁止短褲背心）",
    "注意防曬，七月羅馬非常炎熱",
    "隨身帶水壺補充水分",
])

output_path = "tools/行程網站/Day1_Day2_修正行程.docx"
doc.save(output_path)
print(f"已完成：{output_path}")
