from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x92, 0x46)
    return heading

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_restaurant_table(doc, restaurants, day_name, city, closed_info=None, gambero_info=None):
    add_heading_styled(doc, f"Day {day_name}：{city}", level=2)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['餐廳名稱', '價格', '地址', '電話', '營業時間', '紅蝦評鑑', '當日狀態']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "009246")
    
    for restaurant in restaurants:
        row = table.add_row()
        row.cells[0].text = restaurant['name']
        row.cells[1].text = restaurant.get('price', '—')
        row.cells[2].text = restaurant['address']
        row.cells[3].text = restaurant.get('phone', '-')
        row.cells[4].text = restaurant['hours']
        
        # Gambero Rosso rating
        if gambero_info and restaurant['name'] in gambero_info:
            row.cells[5].text = gambero_info[restaurant['name']]
            for paragraph in row.cells[5].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xCE, 0x2B, 0x37)
                    run.font.bold = True
                    run.font.size = Pt(8)
        else:
            row.cells[5].text = '—'
        
        if closed_info and restaurant['name'] in closed_info:
            row.cells[6].text = closed_info[restaurant['name']]
            for paragraph in row.cells[6].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xCE, 0x2B, 0x37)
                    run.font.bold = True
        else:
            row.cells[6].text = '✅ 營業中'
            for paragraph in row.cells[6].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x92, 0x46)
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(2.5)
    
    doc.add_paragraph()

def add_alternatives_table(doc, alternatives):
    add_heading_styled(doc, "替代方案", level=3)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['替代餐廳', '地址', '營業時間', '特色', '費用/距離']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "CE2B37")
    
    for alt in alternatives:
        row = table.add_row()
        row.cells[0].text = alt['name']
        row.cells[1].text = alt['address']
        row.cells[2].text = alt['hours']
        row.cells[3].text = alt['specialty']
        row.cells[4].text = alt.get('price_distance', '-')
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()

title = doc.add_heading('義大利 12 天親子之旅：餐廳完整指南', 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x00, 0x92, 0x46)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('行程日期：2026 年 7 月 22 日 - 8 月 2 日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

add_heading_styled(doc, '目錄', level=1)
toc_items = [
    'Day 1：羅馬（西班牙階梯 / 許願池區域）',
    'Day 2：羅馬（競技場 / 越河區 Trastevere）',
    'Day 3：羅馬（梵蒂岡 / 鮮花廣場）',
    'Day 4：佛羅倫斯（車站 / 中央市場 / 丁骨牛排）',
    'Day 5：佛羅倫斯（烏菲茲 / 老橋 / 米開朗基羅廣場）',
    'Day 6：比薩 / 佛羅倫斯（米開朗基羅廣場夕陽）',
    'Day 7：威尼斯（聖馬可廣場）',
    'Day 8：威尼斯/彩色島（貢多拉 / 里阿爾托橋）',
    'Day 9：米蘭（大教堂 / 埃馬努埃萊二世拱廊）',
    'Day 10：米蘭（大教堂登頂 / 布雷拉區）',
    'Day 11：米蘭（斯福爾扎城堡 / 購物）',
    'Day 12：米蘭/機場',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

add_heading_styled(doc, '重要提醒', level=1)
reminders = [
    '義大利許多餐廳週日公休（如 Da Enzo al 29、Panzerotti Luini、Osteria Alle Testiere）',
    '午餐尖峰 12:30-14:00，晚餐 20:00 後',
    '熱門餐廳建議提前預約（Da Enzo、Zà Zà、Osteria Alle Testiere 需 3-4 週）',
    '小吃類（Gelato、Panzerotto、Supplì）可隨時購買',
    '營業時間可能因季節變動，建議出發前再確認',
    '紅色標示 = 當日不營業，請選擇替代方案',
    '🦐 標示 = Gambero Rosso 紅蝦評鑑餐廳',
    '費用說明：€ = €10 以下 / €€ = €10-25 / €€€ = €25-50 / €€€€ = €50 以上（每人）',
]
for reminder in reminders:
    doc.add_paragraph(reminder, style='List Bullet')

doc.add_page_break()

# ============== Day 1 ==============
day1_restaurants = [
    {'name': 'Bar del Gambero', 'price': '€€', 'address': 'Via del Gambero, 30', 'phone': '06 679 3102', 'hours': '07:00-20:00'},
    {'name': "McDonald's", 'price': '€', 'address': 'Piazza di Spagna, 46', 'phone': '06 699 23074', 'hours': '08:00-24:00'},
    {'name': 'La Strega Nocciola', 'price': '€', 'address': 'Via della Vite, 100', 'phone': '-', 'hours': '11:00-23:00'},
    {'name': 'Günther Gelato', 'price': '€', 'address': 'Via dei Due Macelli, 108', 'phone': '-', 'hours': '10:00-23:00'},
    {'name': 'Trecaffè', 'price': '€€', 'address': 'Via dei Due Macelli, 107', 'phone': '-', 'hours': '07:00-20:00'},
    {'name': 'Ciampini Roma', 'price': '€€', 'address': 'Piazza di San Lorenzo in Lucina, 29', 'phone': '06 687 7387', 'hours': '07:00-20:00'},
    {'name': 'Pastificio Guerra', 'price': '€', 'address': 'Via della Croce, 8', 'phone': '06 679 3102', 'hours': '週一-六 10:00-19:30\n週日 10:00-16:00'},
    {'name': 'Pinsitaly Trevi', 'price': '€€', 'address': 'Via della Panetteria, 12', 'phone': '06 275 3973', 'hours': '週一-四 10:30-22:30\n週五-日 10:30-23:00'},
    {'name': 'Pane e Salame', 'price': '€€', 'address': 'Via di Santa Maria in Via, 19', 'phone': '-', 'hours': '10:00-20:00'},
]
day1_closed = {}
day1_gambero = {'Pastificio Guerra': '🦐 GR 歷史名店\n1918 年創立，€5'}
add_restaurant_table(doc, day1_restaurants, '1', '羅馬（西班牙階梯 / 許願池區域）', day1_closed, day1_gambero)

day1_alts = [
    {'name': 'Pasta Chef', 'address': 'Via della Croce, 5', 'hours': '10:00-22:00', 'specialty': '站著吃的便宜義大利麵', 'price_distance': '€5 / 步行 1 分鐘'},
    {'name': 'Zia Rosetta', 'address': 'Via della Vite, 98', 'hours': '08:00-20:00', 'specialty': '羅馬式三明治（Rosetta）', 'price_distance': '€4-6 / 步行 2 分鐘'},
    {'name': 'Gelateria del Teatro', 'address': 'Via dei Coronari, 65', 'hours': '12:00-23:00', 'specialty': '手工冰淇淋，口味創意', 'price_distance': '€3-5 / 步行 10 分鐘'},
    {'name': '🦐 Da Mariolino', 'address': "Via Mario de' Fiori, 37", 'hours': '12:00-15:00 & 18:30-23:00', 'specialty': 'GR 2025 專文推薦，羅馬經典創意', 'price_distance': '€20-30 / 西班牙階梯 3 分鐘'},
    {'name': '🦐 Antica Trattoria Al Gallinaccio', 'address': 'Vicolo del Gallinaccio, 6', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'Trevi 區高評價 trattoria', 'price_distance': '€20-30 / 許願池步行 8 分鐘'},
]
add_alternatives_table(doc, day1_alts)

doc.add_page_break()

# ============== Day 2 ==============
day2_restaurants = [
    {'name': "L'Antico Forno di Fontana Trevi", 'price': '€', 'address': 'Via delle Muratte, 90', 'phone': '06 699 24880', 'hours': '07:00-19:00'},
    {'name': '金杯咖啡 (Tazza d\'Oro)', 'price': '€', 'address': 'Via della Panetteria, 44', 'phone': '06 678 9792', 'hours': '07:00-19:00'},
    {'name': 'Pizza Florida', 'price': '€', 'address': 'Via Florida, 25', 'phone': '-', 'hours': '10:00-22:00'},
    {'name': 'Alice Pizza', 'price': '€', 'address': 'Corso Vittorio Emanuele II, 35', 'phone': '-', 'hours': '07:00-22:00'},
    {'name': 'Da Enzo al 29', 'price': '€€', 'address': 'Via dei Vascellari, 29', 'phone': '06 581 2260', 'hours': '週一-六 12:30-15:00 & 18:30-22:30\n週日公休'},
    {'name': 'Venchi', 'price': '€', 'address': 'Piazza della Rotonda, 6', 'phone': '-', 'hours': '10:00-21:00'},
]
day2_closed = {'Da Enzo al 29': '❌ 週日公休'}
day2_gambero = {'Da Enzo al 29': '🦐 Michelin\nBib Gourmand'}
add_restaurant_table(doc, day2_restaurants, '2', '羅馬（競技場 / 越河區 Trastevere）', day2_closed, day2_gambero)

day2_alts = [
    {'name': 'Tonnarello', 'address': 'Via della Paglia, 22, Trastevere', 'hours': '11:00-23:00', 'specialty': '經典羅馬麵，超大份量', 'price_distance': '€12-18 / Trastevere 步行 5 分鐘'},
    {'name': 'Flavio al Velavevodetto', 'address': 'Via di Monte Testaccio, 97', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': '同區頂級碳烤義大利麵', 'price_distance': '€15-25 / Testaccio 步行 10 分鐘'},
    {'name': 'Suppli', 'address': 'Via di San Francesco a Ripa, 143', 'hours': '10:30-21:00', 'specialty': '炸飯糰（Supplì）', 'price_distance': '€2 起 / Trastevere 步行 8 分鐘'},
    {'name': '🦐 Checco er Carettiere', 'address': 'Via Benedetta, 10', 'hours': '12:30-15:00 & 19:00-23:00', 'specialty': 'GR 專文推薦，90 年老店', 'price_distance': '€25-35 / Trastevere 步行 5 分鐘'},
    {'name': '🦐 Trattoria Pennestri', 'address': 'Via Giovanni Da Empoli, 5', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'Michelin Bib Gourmand', 'price_distance': '€20-30 / Ostiense 步行 5 分鐘'},
]
add_alternatives_table(doc, day2_alts)

doc.add_page_break()

# ============== Day 3 ==============
day3_restaurants = [
    {'name': 'Pergamino Caffè', 'price': '€€', 'address': 'Piazza del Risorgimento, 7', 'phone': '06 687 7333', 'hours': '07:00-19:00'},
    {'name': 'Pizza Zizza', 'price': '€', 'address': 'Via delle Fornaci, 11', 'phone': '-', 'hours': '09:00-22:00'},
    {'name': 'Bar Farnese', 'price': '€', 'address': 'Piazza Farnese, 2', 'phone': '-', 'hours': '08:00-20:00'},
    {'name': "Forno Campo de' Fiori", 'price': '€', 'address': "Campo de' Fiori", 'phone': '06 687 8802', 'hours': '07:00-14:00'},
    {'name': 'Da Fortunata', 'price': '€€', 'address': 'Via del Pellegrino, 11', 'phone': '-', 'hours': '12:00-15:00 & 18:00-22:30'},
]
day3_closed = {}
day3_gambero = {'Pizzarium': '🦐 Tre Rotelle\n96 分！'}
add_restaurant_table(doc, day3_restaurants, '3', '羅馬（梵蒂岡 / 鮮花廣場）', day3_closed, day3_gambero)

day3_alts = [
    {'name': '🦐 Pizzarium (Bonci)', 'address': 'Via della Meloria, 43', 'hours': '10:00-23:00', 'specialty': 'GR Tre Rotelle 96分！羅馬最佳', 'price_distance': '€5-8 / 梵蒂岡步行 10 分鐘'},
    {'name': 'Forno Roscioli', 'address': 'Via dei Giubbonari, 21', 'hours': '07:00-20:00', 'specialty': '百年老店白披薩', 'price_distance': '€3-5 / 鮮花廣場步行 8 分鐘'},
    {'name': 'Antico Forno Roscioli', 'address': 'Via dei Chiavari, 34', 'hours': '07:15-19:30', 'specialty': '傳奇披薩店', 'price_distance': '€3-5 / 鮮花廣場步行 8 分鐘'},
]
add_alternatives_table(doc, day3_alts)

doc.add_page_break()

# ============== Day 4 ==============
day4_restaurants = [
    {'name': 'Mercato Centrale Roma', 'price': '€€', 'address': 'Via Giolitti 36', 'phone': '-', 'hours': '08:00-22:00'},
    {'name': 'Trapizzino', 'price': '€', 'address': 'Mercato Centrale Roma 內', 'phone': '-', 'hours': '08:00-21:00'},
    {'name': 'Bonci (Panarium)', 'price': '€', 'address': 'Mercato Centrale Roma 內', 'phone': '-', 'hours': '08:00-21:00'},
    {'name': 'VyTA', 'price': '€€', 'address': 'Roma Termini 車站內', 'phone': '-', 'hours': '06:00-23:00'},
    {'name': 'Gilli', 'price': '€€', 'address': 'Piazza della Repubblica', 'phone': '055 213 896', 'hours': '07:00-02:00'},
    {'name': 'Trattoria Zà Zà', 'price': '€€', 'address': 'Piazza del Mercato Centrale, 26r', 'phone': '055 215 411', 'hours': '11:00-23:00（每天）'},
    {'name': 'Trattoria Dall\'Oste', 'price': '€€€', 'address': 'Via Luigi Alamanni, 3r', 'phone': '-', 'hours': '11:00-23:00'},
    {'name': 'Mercato Centrale Firenze', 'price': '€€', 'address': "Via dell'Ariento", 'phone': '-', 'hours': '10:00-00:00'},
]
day4_closed = {}
day4_gambero = {}
add_restaurant_table(doc, day4_restaurants, '4', '佛羅倫斯（車站 / 中央市場 / 丁骨牛排）', day4_closed, day4_gambero)

day4_alts = [
    {'name': 'Trattoria Mario', 'address': 'Via Rosina, 2', 'hours': '11:30-15:00', 'specialty': '當地人排隊名店', 'price_distance': '€15-25 / 中央市場旁'},
    {'name': "I' Fratellini", 'address': 'Via dei Bambini, 3r', 'hours': '09:00-19:00', 'specialty': '站著喝酒吃三明治', 'price_distance': '€5-8 / 步行 5 分鐘'},
    {'name': 'Panino Divino', 'address': 'Via dei Neri, 51r', 'hours': '09:00-22:00', 'specialty': '平價三明治', 'price_distance': '€5-8 / 步行 8 分鐘'},
    {'name': '🦐 Trattoria Marione', 'address': 'Via della Spada, 27/r', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': '佛羅倫斯經典 trattoria', 'price_distance': '€25-40 / 老橋步行 5 分鐘'},
    {'name': '🦐 Trattoria Alfredo', 'address': 'Via dei Leoni, 14/r', 'hours': '12:00-15:00 & 18:00-22:00', 'specialty': '烏菲茲旁傳統托斯卡尼', 'price_distance': '€20-30 / 烏菲茲步行 2 分鐘'},
]
add_alternatives_table(doc, day4_alts)

doc.add_page_break()

# ============== Day 5 ==============
day5_restaurants = [
    {'name': 'Ben Caffè', 'price': '€€', 'address': 'Via delle Oche, 7', 'phone': '-', 'hours': '07:30-18:00'},
    {'name': "All'Antico Vinaio", 'price': '€', 'address': 'Via dei Neri, 65', 'phone': '055 238 2723', 'hours': '10:00-22:00（每天）'},
    {'name': 'Ditta Artigianale', 'price': '€€', 'address': 'Via dei Neri, 32r', 'phone': '-', 'hours': '07:30-19:00'},
    {'name': 'Rivoire', 'price': '€€€', 'address': 'Piazza della Signoria', 'phone': '055 212 427', 'hours': '07:30-01:00'},
    {'name': 'Trattoria Dall\'Oste', 'price': '€€€', 'address': 'Via Luigi Alamanni, 3r', 'phone': '-', 'hours': '11:00-23:00'},
]
day5_closed = {}
day5_gambero = {"All'Antico Vinaio": '🦐 GR 專文推薦\n50 Top Italy #39'}
add_restaurant_table(doc, day5_restaurants, '5', '佛羅倫斯（烏菲茲 / 老橋 / 米開朗基羅廣場）', day5_closed, day5_gambero)

day5_alts = [
    {'name': 'Mercato Centrale Firenze 二樓', 'address': "Via dell'Ariento", 'hours': '10:00-00:00', 'specialty': '美食街多攤位，適合家庭', 'price_distance': '€8-15 / 車站步行 5 分鐘'},
    {'name': 'Trattoria Sergio Gozzi', 'address': 'Piazza San Lorenzo, 8r', 'hours': '12:00-15:00 & 19:00-22:00', 'specialty': '百年老店，托斯卡尼傳統', 'price_distance': '€20-30 / 中央市場旁'},
    {'name': 'Il Latini', 'address': 'Via dei Palchetti, 6r', 'hours': '12:00-14:30 & 19:00-22:30', 'specialty': '丁骨牛排名店', 'price_distance': '€30-45 / 老橋步行 10 分鐘'},
]
add_alternatives_table(doc, day5_alts)

doc.add_page_break()

# ============== Day 6 ==============
day6_restaurants = [
    {'name': 'Ben Caffè', 'price': '€€', 'address': 'Via delle Oche, 7', 'phone': '-', 'hours': '07:30-18:00'},
    {'name': 'Mercato Centrale Firenze 二樓', 'price': '€€', 'address': "Via dell'Ariento", 'phone': '-', 'hours': '10:00-00:00'},
    {'name': 'Il Tartufo', 'price': '€€', 'address': "Mercato Centrale 二樓", 'phone': '-', 'hours': '10:00-22:00'},
    {'name': 'La Pizzeria', 'price': '€€', 'address': "Mercato Centrale 二樓", 'phone': '-', 'hours': '10:00-22:00'},
    {'name': 'Yellow Bar', 'price': '€€', 'address': 'Via del Proconsolo, 39r', 'phone': '055 211 766', 'hours': '週二-日 12:00-15:00 & 18:45-23:00\n週一公休'},
]
day6_closed = {'Yellow Bar': '❌ 週一公休'}
day6_gambero = {}
add_restaurant_table(doc, day6_restaurants, '6', '比薩 / 佛羅倫斯（米開朗基羅廣場夕陽）', day6_closed, day6_gambero)

day6_alts = [
    {'name': 'Buca Mario', 'address': 'Piazza degli Ottaviani, 16r', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': '1886 年老酒窖餐廳', 'price_distance': '€25-35 / 老橋步行 10 分鐘'},
    {'name': 'Trattoria 4 Leoni', 'address': 'Via dei Vellutini, 1r', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': '松露義大利麵', 'price_distance': '€25-35 / Piazza della Passera'},
    {'name': 'Gelateria della Passera', 'address': 'Via della Passera, 2', 'hours': '12:00-22:00', 'specialty': '手工冰淇淋', 'price_distance': '€3-5 / 步行 2 分鐘'},
    {'name': '🦐 Trattoria Mario', 'address': 'Via Rosina, 2', 'hours': '11:30-15:00', 'specialty': '佛羅倫斯最老 trattoria', 'price_distance': '€15-25 / 中央市場旁'},
]
add_alternatives_table(doc, day6_alts)

doc.add_page_break()

# ============== Day 7 ==============
day7_restaurants = [
    {'name': 'Rosa Salva', 'price': '€€', 'address': 'Piazza San Marco, 87', 'phone': '041 522 7949', 'hours': '每天 08:00-20:00'},
    {'name': 'Osteria da Mariano', 'price': '€€', 'address': 'Corso del Popolo, 201, Mestre', 'phone': '-', 'hours': '12:00-15:00 & 19:00-22:30'},
]
day7_closed = {}
day7_gambero = {'Caffè Florian': "🦐 2 Chicchi\nBar d'Italia 2025"}
add_restaurant_table(doc, day7_restaurants, '7', '威尼斯（聖馬可廣場）', day7_closed, day7_gambero)

day7_alts = [
    {'name': '🦐 Caffè Florian', 'address': 'Piazza San Marco, 57', 'hours': '09:00-01:00', 'specialty': 'GR 2 Chicchi，1720 年歷史', 'price_distance': '€10-15 / 聖馬可廣場'},
    {'name': 'Al Prosecco', 'address': 'Calle dei Fuseri, 6021', 'hours': '10:00-21:00', 'specialty': '平價威尼斯小吃拼盤', 'price_distance': '€8-12 / 步行 3 分鐘'},
    {'name': 'Bacaro Jazz', 'address': 'Calle della Bissa, 5734', 'hours': '18:00-01:00', 'specialty': '站著喝 Spritz + 小食', 'price_distance': '€5-8 / 步行 2 分鐘'},
    {'name': '🦐 Osteria Alle Testiere', 'address': 'Calle del Mondo Novo, 5801', 'hours': '12:00-15:00 & 19:00-22:30\n週日一公休', 'specialty': 'Michelin Plate + OAD #87', 'price_distance': '€€€ / San Marco 步行 3 分鐘\n需提前 3-4 週預約'},
]
add_alternatives_table(doc, day7_alts)

doc.add_page_break()

# ============== Day 8 ==============
day8_restaurants = [
    {'name': 'Fritto Misto', 'price': '€', 'address': '彩色島下船處附近', 'phone': '-', 'hours': '10:00-19:00'},
    {'name': 'Caffè del Doge', 'price': '€€', 'address': 'Calle dei Cinque, 609', 'phone': '041 522 7787', 'hours': '每天 07:00-18:30'},
    {'name': 'Pizzeria da Pino', 'price': '€€', 'address': 'Piazzale Luigi Candiani, 17, Mestre', 'phone': '-', 'hours': '12:00-14:30 & 18:30-23:00'},
    {'name': 'Osteria da Mariano', 'price': '€€', 'address': 'Corso del Popolo, 201, Mestre', 'phone': '-', 'hours': '12:00-15:00 & 19:00-22:30'},
]
day8_closed = {}
day8_gambero = {'Trattoria da Romano': '🦐 GR 推薦\n威尼斯 TOP 10'}
add_restaurant_table(doc, day8_restaurants, '8', '威尼斯/彩色島（貢多拉 / 里阿爾托橋）', day8_closed, day8_gambero)

day8_alts = [
    {'name': 'Osteria alla Madonna', 'address': 'Calle de la Madonna, 594', 'hours': '11:30-23:00', 'specialty': '威尼斯經典海鮮', 'price_distance': '€25-35 / Rialto 步行 3 分鐘'},
    {'name': 'Antiche Carampane', 'address': 'Rio Terà de le Carampane, 1911', 'hours': '12:00-15:00 & 19:00-22:30', 'specialty': '隱藏海鮮名店', 'price_distance': '€30-45 / Rialto 步行 5 分鐘'},
    {'name': 'Dal Zemei', 'address': 'Castello 5801', 'hours': '11:00-22:00', 'specialty': '平價 Fritto Misto 炸海鮮', 'price_distance': '€10-15 / 步行 8 分鐘'},
    {'name': '🦐 Trattoria da Romano', 'address': 'Via Galuppi, 221, Burano', 'hours': '12:00-15:00 & 19:00-22:30', 'specialty': 'GR 推薦，risotto di go', 'price_distance': '€25-40 / Burano 島'},
    {'name': '🦐 Osteria Alle Testiere', 'address': 'Calle del Mondo Novo, 5801', 'hours': '12:00-15:00 & 19:00-22:30\n週日一公休', 'specialty': 'Michelin Plate + OAD #87', 'price_distance': '€€€ / 需提前 3-4 週預約'},
]
add_alternatives_table(doc, day8_alts)

doc.add_page_break()

# ============== Day 9 ==============
day9_restaurants = [
    {'name': 'Orsonero Coffee', 'price': '€€', 'address': 'Via Giuseppe Broggi, 15', 'phone': '-', 'hours': '週一-五 08:00-17:00\n週六日公休'},
    {'name': 'Starbucks Reserve Roastery', 'price': '€€€', 'address': 'Piazza Cordusio 附近', 'phone': '-', 'hours': '08:00-23:00'},
    {'name': 'Osteria del Gambero Rosso', 'price': '€€€', 'address': 'Viale Pasubio, 8', 'phone': '02 657 1208', 'hours': '每天 12:00-15:00 & 19:00-00:00'},
]
day9_closed = {'Orsonero Coffee': '❌ 週六日公休'}
day9_gambero = {'Osteria del Gambero Rosso': '⚠️ 名稱含\nGambero Rosso'}
add_restaurant_table(doc, day9_restaurants, '9', '米蘭（大教堂 / 埃馬努埃萊二世拱廊）', day9_closed, day9_gambero)

day9_alts = [
    {'name': 'Marchesi 1824', 'address': 'Via Monte Napoleone, 9', 'hours': '08:00-19:00', 'specialty': '百年糕點老店', 'price_distance': '€10-15 / 步行 5 分鐘'},
    {'name': 'Cova', 'address': 'Via Montenapoleone, 8', 'hours': '08:00-20:00', 'specialty': '1817 年歷史咖啡館', 'price_distance': '€8-12 / 步行 5 分鐘'},
    {'name': 'Piz', 'address': 'Via Solferino, 16', 'hours': '11:00-15:00 & 18:00-22:30', 'specialty': '米蘭新派切片披薩', 'price_distance': '€8-12 / 步行 10 分鐘'},
    {'name': '🦐 Trattoria del Nuovo Macello', 'address': 'Via Vigevano, 8', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'GR 評鑑最佳 cotoletta', 'price_distance': '€25-35 / Vigevano 站步行 2 分鐘'},
    {'name': '🦐 Osteria del Balabiott', 'address': 'Via Giuseppe Dezza, 24', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'GR 提及，米蘭傳統料理', 'price_distance': '€25-35 / 步行 10 分鐘'},
]
add_alternatives_table(doc, day9_alts)

doc.add_page_break()

# ============== Day 10 ==============
day10_restaurants = [
    {'name': 'Pizzeria Spontini', 'price': '€€', 'address': 'Via Santa Radegonda, 11', 'phone': '-', 'hours': '每天 12:00-15:00 & 18:30-23:00'},
]
day10_closed = {}
day10_gambero = {}
add_restaurant_table(doc, day10_restaurants, '10', '米蘭（大教堂登頂 / 布雷拉區）', day10_closed, day10_gambero)

day10_alts = [
    {'name': 'Dry Milano', 'address': 'Via Solferino, 48', 'hours': '12:00-00:00', 'specialty': '手工披薩 + 精釀啤酒', 'price_distance': '€15-25 / 步行 8 分鐘'},
    {'name': 'Berberè', 'address': 'Via Ciro Menotti, 3', 'hours': '12:00-15:00 & 18:30-23:00', 'specialty': '手工酸種披薩', 'price_distance': '€12-18 / 步行 10 分鐘'},
    {'name': 'Luini', 'address': 'Via Santa Radegonda, 16', 'hours': '10:00-15:00（週日公休）', 'specialty': '€3 起 Panzerotto', 'price_distance': '€4-6 / 大教堂步行 2 分鐘'},
    {'name': '🦐 Trattoria della Gloria', 'address': 'Via Angelo Fumagalli, 1', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'GR 評鑑 trattoria', 'price_distance': '€20-30 / 搭地鐵 10 分鐘'},
    {'name': "🦐 L'Oste Italiano", 'address': 'Via Tommei, 3', 'hours': '週二-六 19:30-22:30\n週日 12:00-15:00', 'specialty': '20+ 種 risotto', 'price_distance': '€20-30 / 搭地鐵 15 分鐘'},
]
add_alternatives_table(doc, day10_alts)

doc.add_page_break()

# ============== Day 11 ==============
day11_restaurants = [
    {'name': 'Panzerotti Luini', 'price': '€', 'address': 'Via Santa Radegonda, 16', 'phone': '02 8646 1917', 'hours': '10:00-20:00（週日公休）'},
    {'name': 'Pizzium', 'price': '€€', 'address': 'Via Andrea Doria, 17', 'phone': '-', 'hours': '12:00-15:00 & 18:00-23:00'},
]
day11_closed = {'Panzerotti Luini': '❌ 週日公休'}
day11_gambero = {'Ratanà': '🦐 GR + Michelin\n最佳 risotto'}
add_restaurant_table(doc, day11_restaurants, '11', '米蘭（斯福爾扎城堡 / 購物）', day11_closed, day11_gambero)

day11_alts = [
    {'name': 'Eataly Milano', 'address': 'Via Merano, 13', 'hours': '10:00-23:00', 'specialty': '大型義大利美食市集', 'price_distance': '€15-25 / 搭地鐵 10 分鐘'},
    {'name': 'Risotto Factory', 'address': 'Via Vigevano, 8', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': '米蘭風味燉飯', 'price_distance': '€15-25 / 步行 10 分鐘'},
    {'name': 'Grom', 'address': 'Via Vigevano, 1', 'hours': '08:00-22:00', 'specialty': '天然手工冰淇淋', 'price_distance': '€3-5 / 步行 10 分鐘'},
    {'name': '🦐 Ratanà', 'address': 'Via Gaetano de Castillia, 28', 'hours': '12:00-15:00 & 19:00-23:00', 'specialty': 'GR + Michelin，最佳 risotto', 'price_distance': '€80/人 / Bosco Verticale 旁'},
    {'name': '🦐 Pavé', 'address': 'Via Felice Casati, 27', 'hours': '07:30-19:30', 'specialty': 'GR 3 Coni Gelato + 最佳 bar', 'price_distance': '€5-10 / 步行 15 分鐘'},
]
add_alternatives_table(doc, day11_alts)

doc.add_page_break()

# ============== Day 12 ==============
day12_restaurants = [
    {'name': 'Pan Brioche', 'price': '€€', 'address': '馬爾彭薩機場內', 'phone': '-', 'hours': '06:00-21:00'},
    {'name': 'Briciole Bar', 'price': '€€', 'address': '馬爾彭薩機場內', 'phone': '-', 'hours': '06:00-21:00'},
    {'name': 'Venchi', 'price': '€', 'address': '馬爾彭薩機場安檢後', 'phone': '-', 'hours': '08:00-20:00'},
]
day12_closed = {}
day12_gambero = {}
add_restaurant_table(doc, day12_restaurants, '12', '米蘭/機場', day12_closed, day12_gambero)

day12_alts = [
    {'name': 'Prada Bar', 'address': 'Milano Centrale 車站', 'hours': '06:00-22:00', 'specialty': '設計感咖啡吧', 'price_distance': '€5-8 / 車站內'},
    {'name': 'Spontini', 'address': 'Piazza Duca d\'Aosta, 8, 車站旁', 'hours': '12:00-15:00 & 18:30-23:00', 'specialty': '米蘭厚片披薩', 'price_distance': '€6-10 / 車站旁'},
    {'name': 'Autogrill', 'address': '機場內', 'hours': '05:00-22:00', 'specialty': '連鎖速食', 'price_distance': '€8-12 / 機場內'},
]
add_alternatives_table(doc, day12_alts)

doc.add_page_break()

# ============== 附錄 ==============
add_heading_styled(doc, '附錄：不營業日總整理', level=1)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Day', '餐廳', '不營業日', '建議替代']
header_row = table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "CE2B37")

closed_summary = [
    {'day': 'Day 2', 'restaurant': 'Da Enzo al 29', 'closed': '週日', 'alt': 'Tonnarello / Flavio al Velavevodetto'},
    {'day': 'Day 6', 'restaurant': 'Yellow Bar', 'closed': '週一', 'alt': 'Buca Mario / Trattoria 4 Leoni'},
    {'day': 'Day 9', 'restaurant': 'Orsonero Coffee', 'closed': '週六日', 'alt': 'Marchesi 1824 / Cova'},
    {'day': 'Day 11', 'restaurant': 'Panzerotti Luini', 'closed': '週日', 'alt': 'Eataly Milano / Risotto Factory'},
]

for item in closed_summary:
    row = table.add_row()
    row.cells[0].text = item['day']
    row.cells[1].text = item['restaurant']
    row.cells[2].text = item['closed']
    row.cells[3].text = item['alt']
    
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph()

# ============== Gambero Rosso 紅蝦評鑑總覽 ==============
doc.add_page_break()
add_heading_styled(doc, 'Gambero Rosso 紅蝦評鑑總覽', level=1)
doc.add_paragraph()

# 評鑑等級說明
add_heading_styled(doc, '評鑑等級說明', level=2)
rating_table = doc.add_table(rows=1, cols=3)
rating_table.style = 'Table Grid'
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['類別', '最高等級', '說明']
header_row = rating_table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "009246")

ratings = [
    {'category': '餐廳', 'level': 'Tre Forchette (★★★)', 'desc': '義大利最高榮譽，共 52 間'},
    {'category': '餐廳', 'level': 'Tre Gamberi (★★)', 'desc': '最佳傳統 trattoria，共 40 間'},
    {'category': '披薩（圓盤）', 'level': 'Tre Spicchi (★★★)', 'desc': '最佳披薩，共 96 間'},
    {'category': '披薩（切片）', 'level': 'Tre Rotelle (★★)', 'desc': '最佳披薩 al taglio，共 16 間'},
    {'category': '咖啡/酒吧', 'level': 'Tre Chicchi (★★★)', 'desc': '最佳咖啡品質'},
    {'category': '麵包', 'level': 'Tre Pani (★★★)', 'desc': '最佳麵包坊，共 67 間'},
    {'category': '冰淇淋', 'level': 'Tre Coni (★★★)', 'desc': '最佳冰淇淋'},
]

for rating in ratings:
    row = rating_table.add_row()
    row.cells[0].text = rating['category']
    row.cells[1].text = rating['level']
    row.cells[2].text = rating['desc']
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# 紅蝦餐廳總結
add_heading_styled(doc, '行程紅蝦餐廳總結', level=2)

summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['天數', '餐廳', '評鑑等級', '費用']
header_row = summary_table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, "CE2B37")

gambero_summary = [
    {'day': 'Day 3', 'restaurant': 'Bonci Pizzarium', 'rating': 'Tre Rotelle (96分)', 'price': '€5-8'},
    {'day': 'Day 5', 'restaurant': "All'Antico Vinaio", 'rating': 'GR 專文推薦', 'price': '€8-12'},
    {'day': 'Day 7', 'restaurant': 'Caffè Florian', 'rating': '2 Chicchi (Bar)', 'price': '€10-15'},
    {'day': 'Day 8', 'restaurant': 'Trattoria da Romano', 'rating': 'GR 推薦', 'price': '€25-40'},
    {'day': 'Day 8', 'restaurant': 'Osteria Alle Testiere', 'rating': 'Michelin Plate + OAD #87', 'price': '€€€'},
    {'day': 'Day 11', 'restaurant': 'Ratanà', 'rating': 'GR + Michelin Guide', 'price': '€80'},
    {'day': 'Day 11', 'restaurant': 'Pavé', 'rating': 'GR 3 Coni + 最佳 bar', 'price': '€5-10'},
]

for item in gambero_summary:
    row = summary_table.add_row()
    row.cells[0].text = item['day']
    row.cells[1].text = item['restaurant']
    row.cells[2].text = item['rating']
    row.cells[3].text = item['price']
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# 注意事項
add_heading_styled(doc, '重要注意事項', level=2)
notes = [
    'Osteria Alle Testiere：需提前 3-4 週預約（僅線上預約），週日一公休',
    'Da Enzo al 29：不接受預約，需現場排隊 30-60 分鐘，週日公休',
    'Ratanà：建議預約，€80/人為較高檔選擇',
    'Gambero Rosso 評鑑更注重傳統在地料理，Michelin 偏向精緻 dining',
    '兩者搭配使用效果最佳',
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('本文件由 AI 自動產生，營業時間可能因季節變動，建議出發前再確認。')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

output_path = r'G:\我的雲端硬碟\2026trip\tools\行程網站\義大利12天餐廳完整指南_v2_GamberoRosso.docx'
doc.save(output_path)
print(f'Word 文件已產生：{output_path}')
